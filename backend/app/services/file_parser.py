"""
Unified file parsing service that handles PDF, DOCX, and plain text files.
"""
import io
from typing import Tuple, Optional
from docx import Document
import docx2txt
import pdfplumber
from app.utils.file_storage import file_storage
from app.utils.file_validation import generate_file_id
from app.utils.text_cleaning import clean_text, normalize_whitespace, remove_encoding_issues


# ============================================================================
# PDF Parser
# ============================================================================

class PDFParser:
    """PDF parsing service using pdfplumber."""
    
    @staticmethod
    def extract_text(pdf_content: bytes, max_pages: Optional[int] = None) -> Tuple[str, Optional[str]]:
        """
        Extract text from PDF content.
        
        Args:
            pdf_content: PDF file content as bytes
            max_pages: Maximum number of pages to process (None for all pages)
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Create file-like object from bytes
            pdf_file = io.BytesIO(pdf_content)
            
            # Extract text from pages
            full_text = []
            
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = len(pdf.pages)
                pages_to_process = min(total_pages, max_pages) if max_pages else total_pages
                
                # Limit pages for very large PDFs to improve performance
                # Process first 50 pages by default, or all pages if less than 50
                if not max_pages and total_pages > 50:
                    pages_to_process = 50
                
                # Iterate through pages
                for page_num in range(1, pages_to_process + 1):
                    try:
                        page = pdf.pages[page_num - 1]
                        # Extract text from page
                        page_text = page.extract_text()
                        
                        if page_text:
                            full_text.append(page_text)
                        
                    except Exception as e:
                        # Log error but continue with other pages
                        error_msg = f"Error extracting text from page {page_num}: {str(e)}"
                        # Continue with other pages
                        continue
            
            # Combine all pages
            combined_text = "\n\n".join(full_text)
            
            if not combined_text or combined_text.strip() == "":
                return "", "No text could be extracted from the PDF"
            
            # Clean and normalize text
            cleaned_text = remove_encoding_issues(combined_text)
            cleaned_text = normalize_whitespace(cleaned_text)
            cleaned_text = clean_text(cleaned_text)
            
            return cleaned_text, None
            
        except Exception as e:
            error_message = f"Error parsing PDF: {str(e)}"
            return "", error_message
    
    @staticmethod
    def extract_text_with_layout(pdf_content: bytes) -> Tuple[str, Optional[str], dict]:
        """
        Extract text from PDF with layout information.
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message, metadata)
        """
        try:
            pdf_file = io.BytesIO(pdf_content)
            
            full_text = []
            metadata = {
                "total_pages": 0,
                "pages_with_text": 0,
                "pages_without_text": 0,
            }
            
            with pdfplumber.open(pdf_file) as pdf:
                metadata["total_pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        # Try to extract text preserving layout
                        # This handles multi-column layouts better
                        page_text = page.extract_text(layout=True)
                        
                        if page_text and page_text.strip():
                            full_text.append(page_text)
                            metadata["pages_with_text"] += 1
                        else:
                            # Fallback to regular extraction
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                full_text.append(page_text)
                                metadata["pages_with_text"] += 1
                            else:
                                metadata["pages_without_text"] += 1
                                
                    except Exception as e:
                        metadata["pages_without_text"] += 1
                        continue
            
            combined_text = "\n\n".join(full_text)
            
            if not combined_text or combined_text.strip() == "":
                return "", "No text could be extracted from the PDF", metadata
            
            # Clean text
            cleaned_text = remove_encoding_issues(combined_text)
            cleaned_text = normalize_whitespace(cleaned_text)
            cleaned_text = clean_text(cleaned_text)
            
            return cleaned_text, None, metadata
            
        except Exception as e:
            error_message = f"Error parsing PDF: {str(e)}"
            return "", error_message, {}
    
    @staticmethod
    def get_pdf_metadata(pdf_content: bytes) -> dict:
        """
        Extract metadata from PDF.
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Dictionary with PDF metadata
        """
        try:
            pdf_file = io.BytesIO(pdf_content)
            
            with pdfplumber.open(pdf_file) as pdf:
                metadata = {
                    "total_pages": len(pdf.pages),
                    "metadata": pdf.metadata or {},
                }
                
                return metadata
                
        except Exception as e:
            return {
                "error": str(e),
                "total_pages": 0,
                "metadata": {}
            }


# ============================================================================
# DOCX Parser
# ============================================================================

class DOCXParser:
    """DOCX parsing service using python-docx and docx2txt."""
    
    @staticmethod
    def extract_text(docx_content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text from DOCX content.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Create file-like object from bytes
            docx_file = io.BytesIO(docx_content)
            
            # Try using python-docx first (better structure preservation)
            try:
                doc = Document(docx_file)
                full_text = []
                
                # Extract text from all paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                combined_text = "\n\n".join(full_text)
                
            except Exception as e:
                # Fallback to docx2txt if python-docx fails
                try:
                    docx_file.seek(0)  # Reset file pointer
                    combined_text = docx2txt.process(docx_file)
                except Exception as e2:
                    error_message = f"Error parsing DOCX with both methods: python-docx ({str(e)}), docx2txt ({str(e2)})"
                    return "", error_message
            
            if not combined_text or combined_text.strip() == "":
                return "", "No text could be extracted from the DOCX file"
            
            # Clean and normalize text
            cleaned_text = remove_encoding_issues(combined_text)
            cleaned_text = normalize_whitespace(cleaned_text)
            cleaned_text = clean_text(cleaned_text)
            
            return cleaned_text, None
            
        except Exception as e:
            error_message = f"Error parsing DOCX: {str(e)}"
            return "", error_message
    
    @staticmethod
    def extract_text_with_structure(docx_content: bytes) -> Tuple[str, Optional[str], dict]:
        """
        Extract text from DOCX with structure information.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message, metadata)
        """
        try:
            docx_file = io.BytesIO(docx_content)
            metadata = {
                "total_paragraphs": 0,
                "total_tables": 0,
                "paragraphs_with_text": 0,
                "extraction_method": "python-docx"
            }
            
            try:
                doc = Document(docx_file)
                
                full_text = []
                paragraphs = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    metadata["total_paragraphs"] += 1
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                        metadata["paragraphs_with_text"] += 1
                
                full_text.extend(paragraphs)
                
                # Extract tables
                metadata["total_tables"] = len(doc.tables)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                combined_text = "\n\n".join(full_text)
                
            except Exception as e:
                # Fallback to docx2txt
                try:
                    docx_file.seek(0)
                    combined_text = docx2txt.process(docx_file)
                    metadata["extraction_method"] = "docx2txt"
                    metadata["total_paragraphs"] = len(combined_text.split('\n'))
                except Exception as e2:
                    error_message = f"Error parsing DOCX: {str(e)}"
                    return "", error_message, metadata
            
            if not combined_text or combined_text.strip() == "":
                return "", "No text could be extracted from the DOCX file", metadata
            
            # Clean text
            cleaned_text = remove_encoding_issues(combined_text)
            cleaned_text = normalize_whitespace(cleaned_text)
            cleaned_text = clean_text(cleaned_text)
            
            return cleaned_text, None, metadata
            
        except Exception as e:
            error_message = f"Error parsing DOCX: {str(e)}"
            return "", error_message, {}
    
    @staticmethod
    def get_docx_metadata(docx_content: bytes) -> dict:
        """
        Extract metadata from DOCX.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Dictionary with DOCX metadata
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            metadata = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "core_properties": {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                }
            }
            
            return metadata
            
        except Exception as e:
            return {
                "error": str(e),
                "total_paragraphs": 0,
                "total_tables": 0,
                "core_properties": {}
            }


# ============================================================================
# Text Input Service
# ============================================================================

class TextInputService:
    """Service for handling plain text input."""
    
    MIN_TEXT_LENGTH = 10  # Minimum text length
    MAX_TEXT_LENGTH = 100000  # Maximum text length (100KB)
    
    @staticmethod
    def validate_text(text: str) -> Tuple[bool, Optional[str]]:
        """
        Validate text input.
        
        Args:
            text: Text to validate
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not text or not isinstance(text, str):
            return False, "Text is required and must be a string"
        
        text_length = len(text.strip())
        
        if text_length < TextInputService.MIN_TEXT_LENGTH:
            return False, f"Text must be at least {TextInputService.MIN_TEXT_LENGTH} characters long"
        
        if text_length > TextInputService.MAX_TEXT_LENGTH:
            return False, f"Text exceeds maximum length of {TextInputService.MAX_TEXT_LENGTH} characters"
        
        return True, None
    
    @staticmethod
    def store_text(
        text: str,
        source_type: str = "resume",
        filename: Optional[str] = None
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Store plain text input in file storage.
        
        Args:
            text: Text content
            source_type: Type of input ('resume' or 'job_description')
            filename: Optional filename for reference
            
        Returns:
            Tuple of (text_id, error_message)
        """
        # Validate source_type
        if source_type not in ["resume", "job_description"]:
            return None, f"Invalid source_type: {source_type}. Must be 'resume' or 'job_description'"
        
        # Validate text
        is_valid, error_message = TextInputService.validate_text(text)
        if not is_valid:
            return None, error_message
        
        # Clean and normalize text
        cleaned_text = normalize_whitespace(text)
        cleaned_text = clean_text(cleaned_text)
        
        # Generate text ID
        text_id = generate_file_id()
        
        # Create filename if not provided
        if not filename:
            filename = f"{source_type}_{text_id[:8]}.txt"
        
        # Convert text to bytes for storage
        text_bytes = cleaned_text.encode('utf-8')
        file_size = len(text_bytes)
        
        # Store in file storage (treating as a text file)
        file_data = file_storage.store_file(
            file_id=text_id,
            filename=filename,
            file_type="txt",
            content=text_bytes,
            file_size=file_size,
            source_type=source_type
        )
        
        # Store cleaned text directly (so it's already "parsed")
        file_storage.update_file_text(text_id, cleaned_text)
        
        return text_id, None
    
    @staticmethod
    def get_text(text_id: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Get stored text by ID.
        
        Args:
            text_id: Text identifier
            
        Returns:
            Tuple of (text, error_message)
        """
        file_data = file_storage.get_file(text_id)
        
        if not file_data:
            return None, "Text not found or session expired"
        
        text = file_data.get("parsed_text")
        
        if not text:
            return None, "Text not available"
        
        return text, None


# ============================================================================
# File Parser Service (Router)
# ============================================================================

class FileParserService:
    """Service for parsing uploaded files."""
    
    def __init__(self):
        """Initialize parsers."""
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()
    
    def parse_file(self, file_id: str) -> Tuple[bool, Optional[str]]:
        """
        Parse a file and store the extracted text.
        
        Args:
            file_id: Unique file identifier
            
        Returns:
            Tuple of (success, error_message)
        """
        # Get file from storage
        file_data = file_storage.get_file(file_id)
        
        if not file_data:
            return False, "File not found or session expired"
        
        # Check if already parsed
        if file_data.get("parsed_text"):
            return True, None
        
        file_type = file_data["file_type"].lower()
        content = file_data["content"]
        
        # Route to appropriate parser
        if file_type == "pdf":
            extracted_text, error = self.pdf_parser.extract_text(content)
            
            if error:
                return False, error
            
            # Store parsed text
            file_storage.update_file_text(file_id, extracted_text)
            return True, None
            
        elif file_type == "docx":
            # Use DOCX parser
            extracted_text, error = self.docx_parser.extract_text(content)
            
            if error:
                return False, error
            
            # Store parsed text
            file_storage.update_file_text(file_id, extracted_text)
            return True, None
            
        elif file_type == "txt":
            # Plain text - just decode and clean
            try:
                text = content.decode('utf-8')
                cleaned_text = normalize_whitespace(text)
                cleaned_text = clean_text(cleaned_text)
                
                file_storage.update_file_text(file_id, cleaned_text)
                return True, None
                
            except Exception as e:
                return False, f"Error parsing text file: {str(e)}"
        
        else:
            return False, f"Unsupported file type: {file_type}"


# ============================================================================
# Global Instances
# ============================================================================

# Parser instances
pdf_parser = PDFParser()
docx_parser = DOCXParser()

# Service instances
file_parser_service = FileParserService()
text_input_service = TextInputService()
